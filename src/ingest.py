# © 2026 Plant Operations Brain Team — ET AI Hackathon. All rights reserved.
# Proprietary & confidential. Unauthorised copying, reuse, or redistribution is prohibited (see LICENSE).
"""
ingest.py
---------
Step 1 of the pipeline: turn messy plant documents (PDF, Excel, Word)
into clean, citable text chunks.

Each chunk remembers WHERE it came from (file name + page) so that later
the AI can cite its source -- this is the foundation of our "trustworthy
answers" differentiator.
"""

from __future__ import annotations
from dataclasses import dataclass, asdict
from pathlib import Path
import fitz  # PyMuPDF


@dataclass
class Chunk:
    """A small, searchable piece of a document that knows its own origin."""
    text: str
    source_file: str
    page: int          # 1-based page number (0 for non-paged formats)
    chunk_id: int

    def as_dict(self) -> dict:
        return asdict(self)


def _split_into_chunks(text: str, max_chars: int = 1100, overlap: int = 150) -> list[str]:
    """
    Break a long page of text into overlapping windows.

    Overlap matters: if an answer sits on the boundary between two chunks,
    the overlap makes sure neither half loses the context.
    """
    text = text.strip()
    if not text:
        return []

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = start + max_chars
        chunks.append(text[start:end])
        start = end - overlap  # step back so windows overlap
        if start < 0:
            start = 0
    return chunks


def _ocr_page(page) -> str:
    """
    OCR a scanned (image-only) PDF page. Two tiers, so scanned inspection forms
    are read even on machines WITHOUT a local OCR binary:

      1. Tesseract (pytesseract) if installed — fast, fully offline.
      2. Vision LLM fallback — needs no local binary, so OCR always works when
         an AI provider is configured (the common demo case).

    Returns "" only if BOTH tiers are unavailable, so the pipeline still
    degrades gracefully instead of crashing.
    """
    try:
        png = page.get_pixmap(dpi=200).tobytes("png")
    except Exception:
        return ""

    # Tier 1 — local Tesseract (offline, cheapest).
    try:
        import pytesseract            # optional dependency
        from PIL import Image
        import io
        txt = pytesseract.image_to_string(Image.open(io.BytesIO(png)))
        if txt and txt.strip():
            return txt
    except Exception:
        pass                          # binary/module missing — fall through

    # Tier 2 — vision LLM (no local binary required).
    try:
        from .llm import vision
        txt = vision(
            "Transcribe ALL text from this scanned industrial document page "
            "verbatim: headings, tables, form fields, equipment tags, dates, "
            "signatures and any legible handwriting. Preserve table structure "
            "with ' | ' between columns. Output ONLY the transcribed text.",
            png, mime="image/png", max_tokens=1500)
        return txt or ""
    except Exception:
        return ""


def load_pdf(path: Path, start_id: int) -> list[Chunk]:
    """
    Read a PDF page-by-page so every chunk keeps its page number.

    If a page has little/no extractable text (a scanned document), we fall
    back to OCR -- this is what lets us ingest old scanned inspection forms,
    a core requirement of "heterogeneous document" ingestion.
    """
    chunks: list[Chunk] = []
    doc = fitz.open(path)
    cid = start_id
    for page_index in range(len(doc)):
        page = doc[page_index]
        page_text = page.get_text("text")
        if len(page_text.strip()) < 20:          # likely a scanned image page
            page_text = _ocr_page(page)
        for piece in _split_into_chunks(page_text):
            chunks.append(
                Chunk(
                    text=piece,
                    source_file=path.name,
                    page=page_index + 1,
                    chunk_id=cid,
                )
            )
            cid += 1
    doc.close()
    return chunks


def load_txt(path: Path, start_id: int) -> list[Chunk]:
    """Plain text / markdown files (handy for synthetic test docs)."""
    text = path.read_text(encoding="utf-8", errors="ignore")
    chunks: list[Chunk] = []
    cid = start_id
    for piece in _split_into_chunks(text):
        chunks.append(Chunk(piece, path.name, 0, cid))
        cid += 1
    return chunks


def load_csv(path: Path, start_id: int) -> list[Chunk]:
    """
    CSV / TSV files (equipment registers, sensor logs, spare-parts lists are
    often exported this way). Rows become 'col | col | col' lines so the table
    stays readable and searchable, with the header repeated in each chunk.
    """
    import csv as _csv

    text = path.read_text(encoding="utf-8", errors="ignore")
    delim = "\t" if (path.suffix.lower() == ".tsv" or "\t" in text[:2000]) else ","
    rows = list(_csv.reader(text.splitlines(), delimiter=delim))
    if not rows:
        return []
    header = " | ".join(str(c) for c in rows[0])
    lines = [f"# {path.stem}", header]
    for r in rows[1:]:
        cells = [str(c) for c in r if str(c).strip()]
        if cells:
            lines.append(" | ".join(cells))
    body = "\n".join(lines)

    chunks, cid = [], start_id
    for piece in _split_into_chunks(body):
        # keep the header on every chunk so rows never lose their column names
        piece = piece if piece.startswith("#") else f"{header}\n{piece}"
        chunks.append(Chunk(piece, path.name, 0, cid))
        cid += 1
    return chunks


def load_xlsx(path: Path, start_id: int) -> list[Chunk]:
    """
    Spreadsheets (maintenance logs, inspection registers are often .xlsx).
    Each sheet becomes text rows; the sheet name is kept as the 'page' label
    via the chunk text header so citations stay meaningful.
    """
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    chunks: list[Chunk] = []
    cid = start_id
    for sheet in wb.worksheets:
        lines = [f"# Sheet: {sheet.title}"]
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append(" | ".join(cells))
        text = "\n".join(lines)
        for piece in _split_into_chunks(text):
            chunks.append(Chunk(piece, path.name, 0, cid))
            cid += 1
    wb.close()
    return chunks


def load_image(path: Path, start_id: int) -> list[Chunk]:
    """
    Engineering drawings / P&IDs / diagrams as image files (.png/.jpg).
    Uses vision (Gemini) to 'read' the drawing -- equipment tags, components,
    connections, valve/instrument labels, notes, title block -- and turns it into
    searchable text so drawings become first-class citizens of the knowledge base.
    """
    import mimetypes
    from .llm import vision

    mime = mimetypes.guess_type(str(path))[0] or "image/png"
    prompt = (
        "This is an industrial engineering drawing / P&ID / technical diagram. "
        "Describe it as clear, readable prose (NOT JSON, NOT bounding boxes). "
        "Cover: every equipment tag (e.g. P-7, V-101) and what it is; how the "
        "components connect / flow between them; all valve and instrument labels; "
        "any dimensions/specifications; all notes; and the title block "
        "(drawing number, title). Be thorough -- this text is the searchable "
        "record of the drawing.\n\n"
        "THEN, on new lines, add a machine-readable section titled exactly "
        "'CONNECTIONS:' listing the process topology, one edge per line in the "
        "format 'TAG_A -> TAG_B : relation' (e.g. 'P-7 -> HX-2 : discharges to'). "
        "Use the real tags from the drawing. This section lets the drawing feed "
        "the knowledge graph, so be precise and include every connection you see."
    )
    try:
        desc = vision(prompt, path.read_bytes(), mime=mime, max_tokens=1500)
    except Exception as e:
        desc = f"[Drawing could not be analysed automatically: {e}]"

    text = f"# Engineering Drawing: {path.name}\n\n{desc}"
    chunks, cid = [], start_id
    for piece in _split_into_chunks(text):
        chunks.append(Chunk(piece, path.name, 0, cid))
        cid += 1
    return chunks


def load_eml(path: Path, start_id: int) -> list[Chunk]:
    """
    Email archives (.eml) -- the problem statement explicitly names 'regulatory
    submissions scattered across email archives'. We keep the headers (who/when/
    subject) because in plants, WHO approved something and WHEN often matters as
    much as the content itself.
    """
    import email
    from email import policy

    msg = email.message_from_bytes(path.read_bytes(), policy=policy.default)
    lines = [
        f"# Email: {msg.get('Subject', '(no subject)')}",
        f"From: {msg.get('From', '?')}",
        f"To: {msg.get('To', '?')}",
        f"Date: {msg.get('Date', '?')}",
        "",
    ]
    body = msg.get_body(preferencelist=("plain", "html"))
    if body is not None:
        content = body.get_content()
        if body.get_content_type() == "text/html":
            # crude but dependency-free HTML -> text
            import re
            content = re.sub(r"<[^>]+>", " ", content)
        lines.append(content)
    # note attachments by name so they're at least discoverable
    for part in msg.iter_attachments():
        lines.append(f"[Attachment: {part.get_filename() or 'unnamed'}]")

    text = "\n".join(lines)
    chunks, cid = [], start_id
    for piece in _split_into_chunks(text):
        chunks.append(Chunk(piece, path.name, 0, cid))
        cid += 1
    return chunks


def load_docx(path: Path, start_id: int) -> list[Chunk]:
    """Word documents (SOPs, procedures, reports are commonly .docx)."""
    from docx import Document

    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Include simple tables too -- procedures often store steps in tables.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    text = "\n".join(paragraphs)
    chunks: list[Chunk] = []
    cid = start_id
    for piece in _split_into_chunks(text):
        chunks.append(Chunk(piece, path.name, 0, cid))
        cid += 1
    return chunks


# Map file extensions to their loader function.
_LOADERS = {
    ".pdf": load_pdf,
    ".txt": load_txt,
    ".md": load_txt,
    ".xlsx": load_xlsx,
    ".xlsm": load_xlsx,
    ".csv": load_csv,
    ".tsv": load_csv,
    ".docx": load_docx,
    ".eml": load_eml,
    ".png": load_image,
    ".jpg": load_image,
    ".jpeg": load_image,
    ".webp": load_image,
}


def ingest_folder(folder: str | Path) -> list[Chunk]:
    """
    Walk a folder, load every supported document, and return all chunks.
    This is the single entry point the rest of the app calls.
    """
    folder = Path(folder)
    all_chunks: list[Chunk] = []
    next_id = 0

    for path in sorted(folder.rglob("*")):
        loader = _LOADERS.get(path.suffix.lower())
        if loader is None:
            continue
        new_chunks = loader(path, next_id)
        all_chunks.extend(new_chunks)
        next_id += len(new_chunks)
        print(f"  ingested {path.name}: {len(new_chunks)} chunks")

    print(f"Total: {len(all_chunks)} chunks from {folder}")
    return all_chunks


if __name__ == "__main__":
    # Quick manual test: python -m src.ingest data/docs
    import sys
    target = sys.argv[1] if len(sys.argv) > 1 else "data/docs"
    ingest_folder(target)
